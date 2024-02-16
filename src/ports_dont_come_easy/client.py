#!/usr/bin/env python3

"""
client.py - update the report document with a custom client title.
"""


import argparse
from datetime import datetime
import sys
import os

from docx import Document


def update_document_title(doc, title_text):
    """
    Update the document title if it exists, or add a new title.
    """

    title_updated = False
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Title':
            paragraph.text = title_text
            title_updated = True
            break

    # If no title paragraph with 'Title' style was found,
    # add a new title at the beginning of the document.
    if not title_updated:
        doc.add_paragraph(title_text, style='Title')


def add_custom_header(doc_path, client_name):
    """
    Add or update the document title with custom information.
    """

    todays_date = datetime.now().strftime('%Y-%m-%d')
    title_text = f"Penetration Testing {todays_date} - Client: {client_name}"

    if os.path.exists(doc_path):
        doc = Document(doc_path)
    else:
        doc = Document()

    update_document_title(doc, title_text)
    doc.save(doc_path)


def main():
    """
    Define the arguments,
    expect a file, create one if missing,
    expect a client name, for the document title,
    call the add_custom_header() function.
    """

    parser = argparse.ArgumentParser(
            description="Update document with custom client title")

    parser.add_argument(
            "-o", "--output",
            help="Output document name (default: pentest-<date>.docx)",
            default=f"pentest-{datetime.now().strftime('%Y-%m-%d')}.docx")

    parser.add_argument(
            "client", help="Client name to include in the document title")

    args = parser.parse_args()

    add_custom_header(args.output, args.client)

    print(f"Document '{args.output}' updated with custom client title.", file=sys.stderr)


if __name__ == "__main__":
    main()
