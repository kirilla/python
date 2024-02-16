#!/usr/bin/env python3

"""
disclaimer.py - adds a disclaimer section to the pentest report
"""


import argparse
from datetime import datetime
import sys

from docx import Document


def add_disclaimer_section(doc_path, disclaimer_text):
    """Add a disclaimer section to the document."""
    try:
        doc = Document(doc_path)
    except FileNotFoundError:
        doc = Document()
        doc.add_heading('Disclaimer', level=1)
    else:
        # Add a page break before the disclaimer for readability
        doc.add_page_break()
        doc.add_heading('Disclaimer', level=1)

    doc.add_paragraph(disclaimer_text)
    doc.save(doc_path)


def main():
    """
    Start the program, define arguments,
    define the disclaimer text,
    and call the add_disclaimer_section() function.
    """

    parser = argparse.ArgumentParser(
            description="Add a disclaimer section to the pentest report")

    parser.add_argument(
            "-o", "--output",
            help="Output document name (default: pentest-<date>.docx)",
            default=f"pentest-{datetime.now().strftime('%Y-%m-%d')}.docx")

    args = parser.parse_args()

    # pylint: disable=line-too-long
    disclaimer_text = "The findings in this report are for authorized use only. The methodologies used are intended for educational purposes and should not be used maliciously."

    add_disclaimer_section(args.output, disclaimer_text)

    print(f"Disclaimer section added to '{args.output}'.", file=sys.stderr)


if __name__ == "__main__":
    main()
