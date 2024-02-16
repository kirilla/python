#!/usr/bin/env python3

"""
intro.py - add or update an introduction section to a pentesting report.
"""


import argparse
from datetime import datetime
import sys
import os

from docx import Document


def add_or_update_intro_section(doc, intro_text):
    """Add or update the Introduction section."""
    # Search for the Introduction heading
    intro_heading_found = False
    for paragraph in doc.paragraphs:
        if paragraph.text == "Introduction" and paragraph.style.name == 'Heading 1':
            intro_heading_found = True
            # Clear existing introduction content
            while paragraph._element.getnext() is not None:
                next_elem = paragraph._element.getnext()
                if next_elem.tag.endswith('heading'):
                    break
                paragraph._element.getparent().remove(next_elem)
            break

    if not intro_heading_found:
        # Add Introduction as the second section
        doc.add_paragraph("Introduction", style='Heading 1')

    # Add the introduction text
    doc.add_paragraph(intro_text)


def main():
    """
    Define arguments, see if we've been given a file work with,
    or whether we should create a new one.
    Call the add_or_update_intro_section()
    to add an intro section to the document.
    """

    parser = argparse.ArgumentParser(
            description="Add or update an Introduction section in the pentest report")

    parser.add_argument(
            "-o", "--output",
            help="Output document name (default: pentest-<current date>.docx)",
            default=f"pentest-{datetime.now().strftime('%Y-%m-%d')}.docx")

    args = parser.parse_args()

    if os.path.exists(args.output):
        doc = Document(args.output)
    else:
        doc = Document()
        doc.add_heading(
                f'Penetration Testing {datetime.now().strftime("%Y-%m-%d")}', 0)
        # Default title

    # pylint: disable=line-too-long
    intro_text = "This document provides an overview of the penetration testing methodology, scope, and objectives."

    add_or_update_intro_section(doc, intro_text)

    doc.save(args.output)

    print(f"Introduction section added/updated in '{args.output}'.", file=sys.stderr)


if __name__ == "__main__":
    main()
