#!/usr/bin/env python3

"""
pinfo.py - a auxilliary script that adds an informative section
to the pentesting port, in the form of a table with common TCP
port numbers and the common services usually hosted on those ports.
"""


import argparse
from datetime import datetime
import sys
import os

from docx import Document


def add_or_update_default_ports_section(doc):
    """
    Add or update the 'Default Ports' section with a table of standard ports.
    """

    section_title = "Default Ports"

    ports_services = [
        (22, "SSH"),
        (25, "SMTP"),
        (53, "DNS"),
        (80, "HTTP"),
        (443, "HTTPS"),
    ]

    # Attempt to find the Default Ports section
    section_found = False
    for paragraph in doc.paragraphs:
        if (paragraph
                and paragraph.text == section_title
                and paragraph.style.name == 'Heading 1'):
            section_found = True
            break

    # If section is found, clear existing table if present
    if section_found:
        for table in doc.tables:
            # Assuming the table follows the section immediately
            p_index = doc.element.body.index(paragraph._element)
            t_index = doc.element.body.index(table._element)
            if t_index > p_index:
                doc.element.body.remove(table._element)
                break

    # If section not found, add it
    if not section_found:
        doc.add_paragraph(section_title, style='Heading 1')

    # Add or update the table
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Port'
    hdr_cells[1].text = 'Service'
    for port, service in ports_services:
        row_cells = table.add_row().cells
        row_cells[0].text = str(port)
        row_cells[1].text = service


def main():
    """
    Define arguments, set up default values,
    see if we have been given a file to work with,
    or if a new one should be created,
    ensure there is a heading,
    call the add_or_update_default_ports_section() function,
    and save the changes.
    """

    todays_date = datetime.now().strftime('%Y-%m-%d')

    parser = argparse.ArgumentParser(
            description="Update or create a document with a 'Default Ports' section")

    parser.add_argument(
            "-o", "--output",
            help="Output document name (default: pentest-<current date>.docx)",
            default=f"pentest-{todays_date}.docx")

    args = parser.parse_args()

    # Check if the document exists to decide on opening or creating a new one.
    if os.path.exists(args.output):
        doc = Document(args.output)
    else:
        doc = Document()
        # Add a default title to ensure the document structure.
        doc.add_heading(f'Penetration Testing {datetime.now().strftime("%Y-%m-%d")}', 0)

    add_or_update_default_ports_section(doc)
    doc.save(args.output)
    print(f"Document '{args.output}' updated with 'Default Ports' section.", file=sys.stderr)

if __name__ == "__main__":
    main()
