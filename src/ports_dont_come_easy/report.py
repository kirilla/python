#!/usr/bin/env python3

"""
report.py - Generate a report from the findings.

This should include open ports, protocol analysis
and vulnerability findings.
"""


import argparse
from datetime import datetime
import sys

from docx import Document


def find_or_create_section(document, section_title):
    """
    Find the paragraph object for a section title or create a new section if not found.
    """
    for paragraph in document.paragraphs:
        if paragraph.text == section_title and paragraph.style == document.styles['Heading 1']:
            return paragraph
    # If section not found, add it at the end of the document
    return document.add_heading(section_title, level=1)


def append_to_section(doc_path, section_title, content):
    """
    Find the document, find the paragraph,
    append the content.
    """

    try:
        doc = Document(doc_path)
    except Exception:
        todays_date = datetime.now().strftime('%Y-%m-%d')
        doc = Document()
        doc.add_heading(f"Penetration Testing {todays_date}", 0)

    # section_paragraph = 
    find_or_create_section(doc, section_title)

    # Insert the content after the section title
    doc.add_paragraph(content)
    doc.save(doc_path)


def main():
    """
    Define arguments, expect a filename, but provide a default name as fallback,
    expect findings on standard input, but handle manual input if no piped input
    is present.
    """

    parser = argparse.ArgumentParser(
            description="Generate a report from findings")

    parser.add_argument(
            "-o", "--output",
            help="Output document name (default: pentest-<current date>.docx)")

    parser.add_argument(
            "-s", "--section", default="Untitled",
            help="Section name under which to append findings (default: Untitled)")

    args = parser.parse_args()

    todays_date = datetime.now().strftime('%Y-%m-%d')

    output_file = args.output if args.output else f"pentest-{todays_date}.docx"

    if not sys.stdin.isatty():
        for line in sys.stdin:
            finding = line.strip()
            if finding:
                append_to_section(output_file, args.section, finding)
                # pylint: disable=line-too-long
                print(f"Appended finding to section '{args.section}' in {output_file}", file=sys.stderr)
    else:
        print("Enter your finding (Ctrl+D to finish):", file=sys.stderr)
        while True:
            try:
                line = input()
                finding = line.strip()
                if finding:
                    append_to_section(output_file, args.section, finding)
                    # pylint: disable=line-too-long
                    print(f"Appended finding to section '{args.section}' in {output_file}", file=sys.stderr)
            except EOFError:
                break


if __name__ == "__main__":
    main()
