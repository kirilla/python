#!/home/kirilla/ws/python/venv/bin/python3
#!/usr/bin/env python3
import argparse
from docx import Document
from datetime import datetime
import platform
import socket
import os


def get_system_info():
    """Gather system meta information."""
    info = {
        "Operating System": platform.system(),
        "OS Version": platform.version(),
        "OS Release": platform.release(),
        "Architecture": platform.machine(),
        "Hostname": socket.gethostname(),
        "IP Address": socket.gethostbyname(socket.gethostname())
    }
    return info


def add_info_to_doc(doc_path, info):
    """Add system information to the specified section in the document."""
    doc = Document(doc_path)

    # Find or add the section title
    section_title = "System Information"
    section_found = False
    for paragraph in doc.paragraphs:
        if paragraph.text == section_title:
            section_found = True
            break
    if not section_found:
        doc.add_heading(section_title, level=1)

    # Add system information
    for key, value in info.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.save(doc_path)


def main():
    parser = argparse.ArgumentParser(description="Add system information to a pentest report")
    parser.add_argument("-o", "--output", help="Output document name (default: pentest-<current date>.docx)")
    args = parser.parse_args()

    output_date = datetime.now().strftime('%Y-%m-%d')
    output_file = args.output if args.output else f"pentest-{output_date}.docx"

    # Add system information to the document
    system_info = get_system_info()
    add_info_to_doc(output_file, system_info)
    print(f"System information added to {output_file}")


if __name__ == "__main__":
    main()
